import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def format_official_document(input_file, output_file):
    # 打开现有文档或创建新文档
    try:
        doc = docx.Document(input_file)
    except:
        doc = docx.Document()
    
    # 设置页面边距（上下左右各2.54厘米）
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # 设置标题格式
    if len(doc.paragraphs) > 0:
        title = doc.paragraphs[0]
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0] if title.runs else title.add_run()
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        title_run.font.size = Pt(16)
        title_run.font.bold = True
    
    # 设置正文格式
    for para in doc.paragraphs[1:]:
        para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        for run in para.runs:
            run.font.name = '仿宋'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
            run.font.size = Pt(16)
            run.font.bold = False
        
        # 设置段落间距和首行缩进
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Pt(32)
    
    # 设置表格格式
    for table in doc.tables:
        # 设置表格居中
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表格内容格式
        for row in table.rows:
            for cell in row.cells:
                # 设置单元格内容格式
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(12)
                        
        # 设置表格边框
        table.style = 'Table Grid'
    
    # 设置图片格式（如果有的话）
    for shape in doc.inline_shapes:
        # 设置图片大小
        if shape.type.value == 3:  # 3表示图片
            # 设置最大宽度为页面宽度的80%
            max_width = Cm(21.0 * 0.8)  # A4纸宽度约21厘米
            
            # 如果图片宽度超过最大宽度，进行缩放
            if shape.width > max_width:
                aspect_ratio = shape.height / shape.width
                shape.width = max_width
                shape.height = int(max_width * aspect_ratio)
            
            # 确保图片高度不超过页面高度
            max_height = Cm(29.7 * 0.8)  # A4纸高度约29.7厘米
            if shape.height > max_height:
                aspect_ratio = shape.width / shape.height
                shape.height = max_height
                shape.width = int(max_height * aspect_ratio)

    # 保存文档
    doc.save(output_file)

# 使用示例
if __name__ == "__main__":
    input_file = "D:/四川省社保卡应用分析报告.docx"  # 输入文件路径
    output_file = "D:/公文格式文档.docx"  # 输出文件路径
    format_official_document(input_file, output_file)